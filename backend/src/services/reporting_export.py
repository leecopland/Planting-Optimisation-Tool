import io
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, RGBColor
from fpdf import FPDF, XPos, YPos

from src.domains.reporting import FarmReportContract

LOGO_PATH_SVG = Path(__file__).parent.parent.parent.parent / "frontend" / "public" / "assets" / "images" / "logo2.svg"
LOGO_PATH_PNG = Path(__file__).parent.parent.parent.parent / "frontend" / "public" / "assets" / "images" / "logo2.png"


def _get_reason_color_docx(result: str) -> RGBColor:
    if any(x in result for x in ["inside", "exact match", "plateau"]):
        return RGBColor(0, 128, 0)
    elif any(x in result for x in ["below", "above"]):
        return RGBColor(220, 0, 0)
    return RGBColor(200, 120, 0)


def _get_reason_color_pdf(result: str) -> tuple:
    if any(x in result for x in ["inside", "exact match", "plateau"]):
        return (0, 128, 0)
    elif any(x in result for x in ["below", "above"]):
        return (220, 0, 0)
    return (200, 120, 0)


def _split_reason(reason: str) -> tuple[str, str]:
    parts = reason.split(":", 1)
    if len(parts) == 2:
        return parts[0].strip(), parts[1].strip()
    return "", reason.strip()


def _add_key_reasons_docx(table, rec):
    if not rec.key_reasons:
        return
    row_cells = table.add_row().cells
    row_cells[0].merge(row_cells[3])
    para = row_cells[0].paragraphs[0]
    for reason in rec.key_reasons:
        factor, result = _split_reason(reason)
        text = f"  {factor}: {result}   " if factor else f"  {result}   "
        run = para.add_run(text)
        run.font.color.rgb = _get_reason_color_docx(result)


def _add_key_reasons_pdf(pdf, rec):
    if not rec.key_reasons:
        return
    for reason in rec.key_reasons:
        factor, result = _split_reason(reason)
        text = f"  {factor}: {result}" if factor else f"  {result}"
        r, g, b = _get_reason_color_pdf(result)
        pdf.set_text_color(r, g, b)
        pdf.cell(0, 6, text, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_text_color(0, 0, 0)


def generate_docx(report: FarmReportContract) -> bytes:
    doc = Document()

    if LOGO_PATH_PNG.exists():
        logo_para = doc.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        logo_para.add_run().add_picture(str(LOGO_PATH_PNG), width=Inches(1.5))

    doc.add_heading("Planting Optimisation Tool - Farm Report", level=1)
    doc.add_paragraph(f"Generated: {report.generated_at.strftime('%Y-%m-%d %H:%M UTC')}")
    doc.add_paragraph("")

    doc.add_heading("Farm Profile", level=2)
    farm = report.farm
    farm_details = [
        ("Farm ID", str(farm.id)),
        ("Farm Owner", farm.user_name),
        ("Rainfall", f"{farm.rainfall_mm} mm"),
        ("Temperature", f"{farm.temperature_celsius} °C"),
        ("Elevation", f"{farm.elevation_m} m"),
        ("pH", str(farm.ph)),
        ("Soil Texture", farm.soil_texture),
        ("Area", f"{farm.area_ha} ha"),
        ("Latitude", str(farm.latitude)),
        ("Longitude", str(farm.longitude)),
    ]

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    header_cells[0].text = "Property"
    header_cells[1].text = "Value"

    for label, value in farm_details:
        row_cells = table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = value

    doc.add_paragraph("")
    doc.add_heading("Species Recommendations", level=2)

    if not report.recommendations:
        doc.add_paragraph("No recommendations available for this farm.")
    else:
        rec_table = doc.add_table(rows=1, cols=4)
        rec_table.style = "Table Grid"
        header_cells = rec_table.rows[0].cells
        header_cells[0].text = "Rank"
        header_cells[1].text = "Species"
        header_cells[2].text = "Common Name"
        header_cells[3].text = "Score"

        for rec in report.recommendations:
            row_cells = rec_table.add_row().cells
            row_cells[0].text = str(rec.rank_overall)
            row_cells[1].text = rec.species_name
            row_cells[2].text = rec.species_common_name
            row_cells[3].text = f"{rec.score_mcda:.2f}"
            _add_key_reasons_docx(rec_table, rec)

    doc.add_paragraph("")
    doc.add_heading("Excluded Species", level=2)

    if not report.exclusions:
        doc.add_paragraph("No species were excluded for this farm.")
    else:
        excl_table = doc.add_table(rows=1, cols=3)
        excl_table.style = "Table Grid"
        excl_header = excl_table.rows[0].cells
        excl_header[0].text = "Species"
        excl_header[1].text = "Common Name"
        excl_header[2].text = "Reasons"

        for excl in report.exclusions:
            row_cells = excl_table.add_row().cells
            row_cells[0].text = excl.species_name
            row_cells[1].text = excl.species_common_name
            para = row_cells[2].paragraphs[0]
            for reason in excl.key_reasons:
                factor, result = _split_reason(reason)
                text = f"{factor}: {result}  " if factor else f"{result}  "
                run = para.add_run(text)
                run.font.color.rgb = RGBColor(220, 0, 0)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def generate_pdf(report: FarmReportContract) -> bytes:
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)

    if LOGO_PATH_SVG.exists():
        pdf.image(str(LOGO_PATH_SVG), x=160, y=8, w=40)

    pdf.set_font("Helvetica", style="B", size=16)
    pdf.cell(0, 10, "Planting Optimisation Tool - Farm Report", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_font("Helvetica", size=10)
    pdf.cell(0, 8, f"Generated: {report.generated_at.strftime('%Y-%m-%d %H:%M UTC')}", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(5)

    pdf.set_font("Helvetica", style="B", size=13)
    pdf.cell(0, 10, "Farm Profile", new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    farm = report.farm
    farm_details = [
        ("Farm ID", str(farm.id)),
        ("Farm Owner", farm.user_name),
        ("Rainfall", f"{farm.rainfall_mm} mm"),
        ("Temperature", f"{farm.temperature_celsius} C"),
        ("Elevation", f"{farm.elevation_m} m"),
        ("pH", str(farm.ph)),
        ("Soil Texture", farm.soil_texture),
        ("Area", f"{farm.area_ha} ha"),
        ("Latitude", str(farm.latitude)),
        ("Longitude", str(farm.longitude)),
    ]

    col_w = 60
    for label, value in farm_details:
        pdf.set_font("Helvetica", style="B", size=10)
        pdf.cell(col_w, 8, label, border=1)
        pdf.set_font("Helvetica", size=10)
        pdf.cell(0, 8, value, border=1, new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    pdf.ln(5)
    pdf.set_font("Helvetica", style="B", size=13)
    pdf.cell(0, 10, "Species Recommendations", new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    if not report.recommendations:
        pdf.set_font("Helvetica", size=10)
        pdf.cell(0, 8, "No recommendations available for this farm.", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    else:
        pdf.set_font("Helvetica", style="B", size=10)
        pdf.cell(15, 8, "Rank", border=1)
        pdf.cell(60, 8, "Species", border=1)
        pdf.cell(60, 8, "Common Name", border=1)
        pdf.cell(0, 8, "Score", border=1, new_x=XPos.LMARGIN, new_y=YPos.NEXT)

        pdf.set_font("Helvetica", size=10)
        for rec in report.recommendations:
            pdf.cell(15, 8, str(rec.rank_overall), border=1)
            pdf.cell(60, 8, rec.species_name[:30], border=1)  # Truncated to fit column width
            pdf.cell(60, 8, rec.species_common_name[:30], border=1)
            pdf.cell(0, 8, f"{rec.score_mcda:.2f}", border=1, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            _add_key_reasons_pdf(pdf, rec)

    pdf.ln(5)
    pdf.set_font("Helvetica", style="B", size=13)
    pdf.cell(0, 10, "Excluded Species", new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    if not report.exclusions:
        pdf.set_font("Helvetica", size=10)
        pdf.cell(0, 8, "No species were excluded for this farm.", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    else:
        pdf.set_font("Helvetica", style="B", size=10)
        pdf.cell(60, 8, "Species", border=1)
        pdf.cell(60, 8, "Common Name", border=1)
        pdf.cell(0, 8, "Reason", border=1, new_x=XPos.LMARGIN, new_y=YPos.NEXT)

        for excl in report.exclusions:
            pdf.set_font("Helvetica", size=10)
            pdf.cell(60, 8, excl.species_name[:30], border=1)
            pdf.cell(60, 8, excl.species_common_name[:30], border=1)
            pdf.cell(0, 8, "", border=1, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            _add_key_reasons_pdf(pdf, excl)

    return bytes(pdf.output())


def generate_all_farms_docx(reports: list[FarmReportContract]) -> bytes:
    doc = Document()

    doc.add_heading("Planting Optimisation Tool - All Farms Report", level=1)
    if reports:
        doc.add_paragraph(f"Generated: {reports[0].generated_at.strftime('%Y-%m-%d %H:%M UTC')}")
    doc.add_paragraph(f"Total farms: {len(reports)}")
    doc.add_paragraph("")

    for report in reports:
        farm = report.farm

        if LOGO_PATH_PNG.exists():
            logo_para = doc.add_paragraph()
            logo_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            logo_para.add_run().add_picture(str(LOGO_PATH_PNG), width=Inches(1.5))

        doc.add_heading(f"Farm {farm.id}", level=2)
        doc.add_heading("Farm Profile", level=3)
        farm_details = [
            ("Farm ID", str(farm.id)),
            ("Farm Owner", farm.user_name),
            ("Rainfall", f"{farm.rainfall_mm} mm"),
            ("Temperature", f"{farm.temperature_celsius} °C"),
            ("Elevation", f"{farm.elevation_m} m"),
            ("pH", str(farm.ph)),
            ("Soil Texture", farm.soil_texture),
            ("Area", f"{farm.area_ha} ha"),
            ("Latitude", str(farm.latitude)),
            ("Longitude", str(farm.longitude)),
        ]

        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        header_cells = table.rows[0].cells
        header_cells[0].text = "Property"
        header_cells[1].text = "Value"

        for label, value in farm_details:
            row_cells = table.add_row().cells
            row_cells[0].text = label
            row_cells[1].text = value

        doc.add_paragraph("")
        doc.add_heading("Species Recommendations", level=3)

        if not report.recommendations:
            doc.add_paragraph("No recommendations available for this farm.")
        else:
            rec_table = doc.add_table(rows=1, cols=4)
            rec_table.style = "Table Grid"
            rec_header = rec_table.rows[0].cells
            rec_header[0].text = "Rank"
            rec_header[1].text = "Species"
            rec_header[2].text = "Common Name"
            rec_header[3].text = "Score"

            for rec in report.recommendations:
                row_cells = rec_table.add_row().cells
                row_cells[0].text = str(rec.rank_overall)
                row_cells[1].text = rec.species_name
                row_cells[2].text = rec.species_common_name
                row_cells[3].text = f"{rec.score_mcda:.2f}"
                _add_key_reasons_docx(rec_table, rec)

        doc.add_paragraph("")
        doc.add_heading("Excluded Species", level=3)

        if not report.exclusions:
            doc.add_paragraph("No species were excluded for this farm.")
        else:
            excl_table = doc.add_table(rows=1, cols=3)
            excl_table.style = "Table Grid"
            excl_header = excl_table.rows[0].cells
            excl_header[0].text = "Species"
            excl_header[1].text = "Common Name"
            excl_header[2].text = "Reasons"

            for excl in report.exclusions:
                row_cells = excl_table.add_row().cells
                row_cells[0].text = excl.species_name
                row_cells[1].text = excl.species_common_name
                para = row_cells[2].paragraphs[0]
                for reason in excl.key_reasons:
                    factor, result = _split_reason(reason)
                    text = f"{factor}: {result}  " if factor else f"{result}  "
                    run = para.add_run(text)
                    run.font.color.rgb = RGBColor(220, 0, 0)

        if report != reports[-1]:
            doc.add_page_break()

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def generate_all_farms_pdf(reports: list[FarmReportContract]) -> bytes:
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)

    for index, report in enumerate(reports):
        pdf.add_page()
        farm = report.farm

        if LOGO_PATH_SVG.exists():
            pdf.image(str(LOGO_PATH_SVG), x=160, y=8, w=40)

        pdf.set_font("Helvetica", style="B", size=16)
        pdf.cell(0, 10, f"Farm {farm.id} - Planting Report", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_font("Helvetica", size=10)
        pdf.cell(0, 8, f"Generated: {report.generated_at.strftime('%Y-%m-%d %H:%M UTC')}", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.cell(0, 8, f"Farm {index + 1} of {len(reports)}", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(5)

        pdf.set_font("Helvetica", style="B", size=13)
        pdf.cell(0, 10, "Farm Profile", new_x=XPos.LMARGIN, new_y=YPos.NEXT)

        farm_details = [
            ("Farm ID", str(farm.id)),
            ("Farm Owner", farm.user_name),
            ("Rainfall", f"{farm.rainfall_mm} mm"),
            ("Temperature", f"{farm.temperature_celsius} C"),
            ("Elevation", f"{farm.elevation_m} m"),
            ("pH", str(farm.ph)),
            ("Soil Texture", farm.soil_texture),
            ("Area", f"{farm.area_ha} ha"),
            ("Latitude", str(farm.latitude)),
            ("Longitude", str(farm.longitude)),
        ]

        col_w = 60
        for label, value in farm_details:
            pdf.set_font("Helvetica", style="B", size=10)
            pdf.cell(col_w, 8, label, border=1)
            pdf.set_font("Helvetica", size=10)
            pdf.cell(0, 8, value, border=1, new_x=XPos.LMARGIN, new_y=YPos.NEXT)

        pdf.ln(5)
        pdf.set_font("Helvetica", style="B", size=13)
        pdf.cell(0, 10, "Species Recommendations", new_x=XPos.LMARGIN, new_y=YPos.NEXT)

        if not report.recommendations:
            pdf.set_font("Helvetica", size=10)
            pdf.cell(0, 8, "No recommendations available for this farm.", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        else:
            pdf.set_font("Helvetica", style="B", size=10)
            pdf.cell(15, 8, "Rank", border=1)
            pdf.cell(60, 8, "Species", border=1)
            pdf.cell(60, 8, "Common Name", border=1)
            pdf.cell(0, 8, "Score", border=1, new_x=XPos.LMARGIN, new_y=YPos.NEXT)

            pdf.set_font("Helvetica", size=10)
            for rec in report.recommendations:
                pdf.cell(15, 8, str(rec.rank_overall), border=1)
                pdf.cell(60, 8, rec.species_name[:30], border=1)  # Truncated to fit column width
                pdf.cell(60, 8, rec.species_common_name[:30], border=1)
                pdf.cell(0, 8, f"{rec.score_mcda:.2f}", border=1, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                _add_key_reasons_pdf(pdf, rec)

        pdf.ln(5)
        pdf.set_font("Helvetica", style="B", size=13)
        pdf.cell(0, 10, "Excluded Species", new_x=XPos.LMARGIN, new_y=YPos.NEXT)

        if not report.exclusions:
            pdf.set_font("Helvetica", size=10)
            pdf.cell(0, 8, "No species were excluded for this farm.", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        else:
            pdf.set_font("Helvetica", style="B", size=10)
            pdf.cell(60, 8, "Species", border=1)
            pdf.cell(60, 8, "Common Name", border=1)
            pdf.cell(0, 8, "Reason", border=1, new_x=XPos.LMARGIN, new_y=YPos.NEXT)

            for excl in report.exclusions:
                pdf.set_font("Helvetica", size=10)
                pdf.cell(60, 8, excl.species_name[:30], border=1)
                pdf.cell(60, 8, excl.species_common_name[:30], border=1)
                pdf.cell(0, 8, "", border=1, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                _add_key_reasons_pdf(pdf, excl)

    return bytes(pdf.output())
